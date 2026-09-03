import logging
import re
import uuid
from functools import lru_cache
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime

import chromadb
from chromadb.config import Settings as ChromaSettings
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings

from backend.config import settings

logger = logging.getLogger("agentic_rag.rag")


def _make_metadata(source: str, page: int = None, extra: Dict[str, Any] = None) -> Dict[str, Any]:
    meta = {"source": source}
    if page is not None:
        meta["page"] = page
    if extra:
        meta.update(extra)
    return meta


def load_pdf(file_path: Path) -> List[Document]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        documents = []

        logger.info(f"Loading PDF: {file_path.name} ({len(reader.pages)} pages)")

        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if text and text.strip():
                doc = Document(
                    page_content=text.strip(),
                    metadata=_make_metadata(
                        source=file_path.name,
                        page=page_num,
                        extra={"file_type": "pdf", "total_pages": len(reader.pages)}
                    )
                )
                documents.append(doc)

        if not documents:
            raise ValueError(f"PDF '{file_path.name}' contains no extractable text")

        logger.info(f"PDF loaded: {len(documents)} pages with text extracted")
        return documents

    except Exception as e:
        logger.error(f"Failed to load PDF {file_path.name}: {e}")
        raise ValueError(f"Could not load PDF '{file_path.name}': {e}")


def load_txt(file_path: Path) -> List[Document]:
    try:
        try:
            text = file_path.read_text(encoding="utf-8-sig")
        except UnicodeDecodeError:
            text = file_path.read_text(encoding="latin-1")

        text = text.strip()
        if not text:
            raise ValueError("Text file is empty")

        logger.info(f"Loading TXT: {file_path.name} ({len(text)} characters)")

        doc = Document(
            page_content=text,
            metadata=_make_metadata(
                source=file_path.name,
                extra={"file_type": "txt", "char_count": len(text)}
            )
        )

        logger.info("TXT loaded successfully")
        return [doc]

    except ValueError:
        raise
    except Exception as e:
        logger.error(f"Failed to load TXT {file_path.name}: {e}")
        raise ValueError(f"Could not load TXT '{file_path.name}': {e}")


def load_docx(file_path: Path) -> List[Document]:
    try:
        # pyrefly: ignore [missing-import]
        from docx import Document as DocxDocument

        doc_obj = DocxDocument(str(file_path))
        paragraphs = [p.text.strip() for p in doc_obj.paragraphs if p.text.strip()]

        for table in doc_obj.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        if not paragraphs:
            raise ValueError("DOCX file appears to be empty")

        full_text = "\n\n".join(paragraphs)
        logger.info(f"Loading DOCX: {file_path.name} ({len(paragraphs)} paragraphs)")

        doc = Document(
            page_content=full_text,
            metadata=_make_metadata(
                source=file_path.name,
                extra={
                    "file_type": "docx",
                    "paragraph_count": len(paragraphs)
                }
            )
        )

        logger.info("DOCX loaded successfully")
        return [doc]

    except ValueError:
        raise
    except Exception as e:
        logger.error(f"Failed to load DOCX {file_path.name}: {e}")
        raise ValueError(f"Could not load DOCX '{file_path.name}': {e}")


def load_document(file_path: Path) -> List[Document]:
    ext = file_path.suffix.lower()
    loaders = {
        ".pdf": load_pdf,
        ".txt": load_txt,
        ".docx": load_docx,
    }

    if ext not in loaders:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Supported types: {', '.join(loaders.keys())}"
        )

    return loaders[ext](file_path)


def create_splitter(
    chunk_size: int = None,
    chunk_overlap: int = None
) -> RecursiveCharacterTextSplitter:
    chunk_size = settings.chunk_size if chunk_size is None else chunk_size
    chunk_overlap = settings.chunk_overlap if chunk_overlap is None else chunk_overlap

    return RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", "! ", "? ", "; ", ", ", " ", ""],
        length_function=len,
        is_separator_regex=False
    )


def split_documents(
    documents: List[Document],
    chunk_size: int = None,
    chunk_overlap: int = None
) -> List[Document]:
    splitter = create_splitter(chunk_size, chunk_overlap)
    chunks = splitter.split_documents(documents)

    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_index"] = i

    logger.info(
        f"Chunking complete: {len(documents)} documents → "
        f"{len(chunks)} chunks "
        f"(size={chunk_size or settings.chunk_size}, "
        f"overlap={chunk_overlap or settings.chunk_overlap})"
    )
    return chunks


@lru_cache(maxsize=1)
def get_embeddings() -> HuggingFaceEmbeddings:
    logger.info(f"Loading embedding model: {settings.embedding_model}")
    embeddings = HuggingFaceEmbeddings(
        model_name=settings.embedding_model,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True}
    )
    logger.info(f"Embedding model loaded successfully: {settings.embedding_model}")
    return embeddings


class VectorStore:
    def __init__(self):
        self._client = None
        self._collection = None
        self._initialized = False

    def _ensure_initialized(self):
        if self._initialized:
            return

        try:
            self._client = chromadb.PersistentClient(
                path=str(settings.chroma_dir),
                settings=ChromaSettings(
                    anonymized_telemetry=False,
                    allow_reset=True
                )
            )

            self._collection = self._client.get_or_create_collection(
                name=settings.chroma_collection_name,
                metadata={"hnsw:space": "cosine"}
            )

            self._initialized = True
            doc_count = self._collection.count()
            logger.info(
                f"ChromaDB connected: collection='{settings.chroma_collection_name}', "
                f"existing_chunks={doc_count}"
            )

        except Exception as e:
            logger.error(f"ChromaDB initialization failed: {e}")
            raise RuntimeError(f"Could not connect to ChromaDB: {e}")

    @property
    def collection(self):
        self._ensure_initialized()
        return self._collection

    def add_documents(self, chunks: List[Document]) -> int:
        self._ensure_initialized()

        if not chunks:
            logger.warning("add_documents called with empty chunk list")
            return 0

        try:
            embeddings_model = get_embeddings()
            texts = [chunk.page_content for chunk in chunks]
            metadatas = [chunk.metadata for chunk in chunks]

            ids = []
            for chunk in chunks:
                source = chunk.metadata.get("source", "unknown")
                chunk_idx = chunk.metadata.get("chunk_index", 0)
                unique_id = f"{source}_{chunk_idx}_{uuid.uuid4().hex[:8]}"
                ids.append(unique_id)

            logger.info(f"Generating embeddings for {len(texts)} chunks...")
            embeddings_list = embeddings_model.embed_documents(texts)

            self._collection.add(
                ids=ids,
                embeddings=embeddings_list,
                documents=texts,
                metadatas=metadatas
            )

            logger.info(f"Successfully added {len(chunks)} chunks to ChromaDB")
            return len(chunks)

        except Exception as e:
            logger.error(f"Failed to add documents to ChromaDB: {e}")
            raise RuntimeError(f"Vector store add failed: {e}")

    def search(
        self,
        query: str,
        top_k: int = None,
        filter_source: Optional[str] = None
    ) -> List[Tuple[Document, float]]:
        self._ensure_initialized()

        if self._collection.count() == 0:
            logger.warning("Search attempted but vector store is empty")
            return []

        top_k = top_k or settings.top_k

        try:
            embeddings_model = get_embeddings()
            query_embedding = embeddings_model.embed_query(query)

            where = None
            if filter_source:
                where = {"source": {"$eq": filter_source}}

            results = self._collection.query(
                query_embeddings=[query_embedding],
                n_results=min(top_k, self._collection.count()),
                where=where,
                include=["documents", "metadatas", "distances"]
            )

            output = []
            if results["documents"] and results["documents"][0]:
                for text, meta, distance in zip(
                    results["documents"][0],
                    results["metadatas"][0],
                    results["distances"][0]
                ):
                    score = 1.0 - (distance / 2.0)
                    doc = Document(page_content=text, metadata=meta)
                    output.append((doc, score))

            logger.info(
                f"Search returned {len(output)} results for "
                f"query='{query[:50]}...' (filter={filter_source})"
            )
            return output

        except Exception as e:
            logger.error(f"ChromaDB search failed: {e}")
            raise RuntimeError(f"Vector search failed: {e}")

    def delete_document(self, source_name: str) -> int:
        self._ensure_initialized()

        try:
            results = self._collection.get(
                where={"source": {"$eq": source_name}},
                include=["documents"]
            )

            if not results["ids"]:
                logger.warning(f"No chunks found for document: {source_name}")
                return 0

            chunk_count = len(results["ids"])
            self._collection.delete(ids=results["ids"])
            logger.info(f"Deleted {chunk_count} chunks for document: {source_name}")
            return chunk_count

        except Exception as e:
            logger.error(f"Failed to delete document {source_name}: {e}")
            raise RuntimeError(f"Could not delete document '{source_name}': {e}")

    def list_indexed_documents(self) -> List[dict]:
        self._ensure_initialized()

        try:
            if self._collection.count() == 0:
                return []

            results = self._collection.get(include=["metadatas"])
            source_counts = {}
            for meta in results["metadatas"]:
                source = meta.get("source", "unknown")
                source_counts[source] = source_counts.get(source, 0) + 1

            return [
                {"source": source, "chunk_count": count}
                for source, count in source_counts.items()
            ]

        except Exception as e:
            logger.error(f"Failed to list indexed documents: {e}")
            return []

    def get_total_chunk_count(self) -> int:
        self._ensure_initialized()
        return self._collection.count()

    def is_document_indexed(self, source_name: str) -> bool:
        indexed = self.list_indexed_documents()
        return any(d["source"] == source_name for d in indexed)

    def get_chunk_count_for_document(self, source_name: str) -> int:
        indexed = self.list_indexed_documents()
        for d in indexed:
            if d["source"] == source_name:
                return d["chunk_count"]
        return 0

    def get_document_chunks(self, source_name: str) -> dict:
        self._ensure_initialized()
        return self._collection.get(
            where={"source": {"$eq": source_name}},
            include=["documents", "metadatas"]
        )


_vector_store_instance: Optional[VectorStore] = None


def get_vector_store() -> VectorStore:
    global _vector_store_instance
    if _vector_store_instance is None:
        _vector_store_instance = VectorStore()
    return _vector_store_instance


@dataclass
class RetrievedChunk:
    content: str
    source: str
    page: Optional[int]
    score: float
    chunk_index: int


def retrieve(
    query: str,
    top_k: int = None,
    filter_source: Optional[str] = None,
    min_score: float = 0.0
) -> List[RetrievedChunk]:
    vector_store = get_vector_store()
    top_k = top_k or settings.top_k

    raw_results = vector_store.search(
        query=query,
        top_k=top_k,
        filter_source=filter_source
    )

    chunks = []
    for doc, score in raw_results:
        if score < min_score:
            continue

        chunk = RetrievedChunk(
            content=doc.page_content,
            source=doc.metadata.get("source", "unknown"),
            page=doc.metadata.get("page", None),
            score=score,
            chunk_index=doc.metadata.get("chunk_index", 0)
        )
        chunks.append(chunk)

    logger.info(
        f"Retrieved {len(chunks)} chunks for query: '{query[:60]}' "
        f"(min_score={min_score}, filter={filter_source})"
    )
    return chunks


def format_chunks_for_context(chunks: List[RetrievedChunk]) -> str:
    if not chunks:
        return "No relevant information found in the uploaded documents."

    formatted_parts = []
    for i, chunk in enumerate(chunks, start=1):
        source_info = f"[Source: {chunk.source}"
        if chunk.page:
            source_info += f", Page {chunk.page}"
        source_info += "]"

        formatted_parts.append(
            f"--- Context Chunk {i} {source_info} ---\n{chunk.content}"
        )

    return "\n\n".join(formatted_parts)


def sanitize_filename(filename: str) -> str:
    filename = Path(filename).name
    filename = re.sub(r'[^\w\-_\.]', '_', filename)
    filename = filename.lstrip('.')
    if not filename:
        filename = "document"
    return filename


def validate_file(filename: str, file_size_bytes: int) -> Tuple[bool, str]:
    ext = Path(filename).suffix.lower().lstrip(".")
    if ext not in settings.allowed_extensions_list:
        return False, (
            f"File type '.{ext}' is not supported. "
            f"Allowed types: {', '.join(settings.allowed_extensions_list)}"
        )

    max_bytes = settings.max_upload_size_mb * 1024 * 1024
    if file_size_bytes > max_bytes:
        return False, (
            f"File size ({file_size_bytes / 1024 / 1024:.1f} MB) exceeds "
            f"the maximum limit of {settings.max_upload_size_mb} MB"
        )

    return True, ""


_document_registry: dict = {}


def get_document_registry() -> dict:
    return _document_registry


def register_document(filename: str, size_bytes: int, status: str = "pending") -> dict:
    _document_registry[filename] = {
        "name": filename,
        "size_bytes": size_bytes,
        "uploaded_at": datetime.now().isoformat(),
        "status": status,
        "chunk_count": 0,
        "error_message": None
    }
    return _document_registry[filename]


def update_document_status(
    filename: str,
    status: str,
    chunk_count: int = None,
    error_message: str = None
) -> None:
    if filename in _document_registry:
        _document_registry[filename]["status"] = status
        if chunk_count is not None:
            _document_registry[filename]["chunk_count"] = chunk_count
        if error_message is not None:
            _document_registry[filename]["error_message"] = error_message


def get_all_documents() -> list:
    vector_store = get_vector_store()
    indexed_docs = {
        d["source"]: d["chunk_count"]
        for d in vector_store.list_indexed_documents()
    }

    data_dir = settings.data_dir
    result = []
    seen = set()

    for name, info in _document_registry.items():
        seen.add(name)
        chunk_count = indexed_docs.get(name, info.get("chunk_count", 0))
        result.append({
            **info,
            "chunk_count": chunk_count,
            "status": "ready" if name in indexed_docs else info.get("status", "unknown")
        })

    for source, chunk_count in indexed_docs.items():
        if source not in seen:
            seen.add(source)
            file_path = data_dir / source
            result.append({
                "name": source,
                "size_bytes": file_path.stat().st_size if file_path.exists() else 0,
                "uploaded_at": "Unknown",
                "status": "ready",
                "chunk_count": chunk_count,
                "error_message": None
            })

    return result


async def ingest_document(file_content: bytes, original_filename: str) -> dict:
    safe_filename = sanitize_filename(original_filename)
    logger.info(f"Starting ingestion: {safe_filename} ({len(file_content)} bytes)")

    is_valid, error_msg = validate_file(safe_filename, len(file_content))
    if not is_valid:
        logger.warning(f"File validation failed: {error_msg}")
        return {"success": False, "error": error_msg}

    register_document(safe_filename, len(file_content), status="processing")
    file_path = settings.data_dir / safe_filename

    try:
        file_path.write_bytes(file_content)
        logger.info(f"File saved: {file_path}")
    except Exception as e:
        error_msg = f"Failed to save file: {e}"
        update_document_status(safe_filename, "error", error_message=error_msg)
        return {"success": False, "error": error_msg}

    try:
        logger.info(f"Loading document: {safe_filename}")
        raw_documents = load_document(file_path)
        logger.info(f"Loaded {len(raw_documents)} document sections")
    except Exception as e:
        error_msg = f"Failed to load document: {e}"
        logger.error(error_msg)
        update_document_status(safe_filename, "error", error_message=error_msg)
        return {"success": False, "error": error_msg}

    try:
        logger.info(f"Chunking document: {safe_filename}")
        chunks = split_documents(raw_documents)
        if not chunks:
            error_msg = "Document appears to be empty or has no extractable text"
            update_document_status(safe_filename, "error", error_message=error_msg)
            return {"success": False, "error": error_msg}
        logger.info(f"Created {len(chunks)} chunks")
    except Exception as e:
        error_msg = f"Failed to chunk document: {e}"
        logger.error(error_msg)
        update_document_status(safe_filename, "error", error_message=error_msg)
        return {"success": False, "error": error_msg}

    try:
        logger.info(f"Generating embeddings and indexing: {safe_filename}")
        vector_store = get_vector_store()
        if vector_store.is_document_indexed(safe_filename):
            logger.info(f"Re-indexing document: removing old chunks for {safe_filename}")
            vector_store.delete_document(safe_filename)

        added_count = vector_store.add_documents(chunks)
        update_document_status(safe_filename, "ready", chunk_count=added_count)

        logger.info(
            f"Ingestion complete: {safe_filename} | "
            f"{len(raw_documents)} sections → {added_count} chunks indexed"
        )
        return {
            "success": True,
            "document": _document_registry[safe_filename],
            "message": (
                f"Successfully indexed '{safe_filename}': "
                f"{added_count} chunks stored in vector database"
            )
        }
    except Exception as e:
        error_msg = f"Failed to index document embeddings: {e}"
        logger.error(error_msg)
        update_document_status(safe_filename, "error", error_message=error_msg)
        return {"success": False, "error": error_msg}


async def delete_document(filename: str) -> dict:
    logger.info(f"Deleting document: {filename}")
    errors = []

    try:
        vector_store = get_vector_store()
        deleted_chunks = vector_store.delete_document(filename)
        logger.info(f"Removed {deleted_chunks} chunks from vector store")
    except Exception as e:
        errors.append(f"Vector store error: {e}")
        logger.error(f"Failed to delete from vector store: {e}")

    try:
        file_path = settings.data_dir / filename
        if file_path.exists():
            file_path.unlink()
            logger.info(f"File deleted from disk: {file_path}")
    except Exception as e:
        errors.append(f"File deletion error: {e}")
        logger.error(f"Failed to delete file from disk: {e}")

    if filename in _document_registry:
        del _document_registry[filename]

    if errors:
        return {
            "success": False,
            "message": f"Partial deletion with errors: {'; '.join(errors)}"
        }

    return {
        "success": True,
        "message": f"Document '{filename}' deleted successfully"
    }
